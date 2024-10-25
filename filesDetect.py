import requests  # type: ignore
import json
import os
import sys
import base64
import time
import re
import docx
from datetime import datetime
from pathlib import Path

BEARER_TOKEN = None

# get full file path to current script
script_path = Path(__file__).resolve()
script_directory = script_path.parent
script_name = script_path.name
script_directory = str(script_directory).replace(script_name, '')
current_directory = script_directory

file_path = current_directory + '/' + 'detect_params.json'  #get params file
with open(file_path, 'r') as file:
    params = json.load(file)

SKYFLOW_ACCOUNT_ID  = params['AccountId']
VAULT_ID            = params['Vault_ID']
VAULT_URL           = params['Vault_URL']
DATA_TYPES          = params['Data_Types']      #detect supported file types
MAX_ATTEMPTS        = params['Max_Api_Attempts']
FULL_REDACTION = params['fullRedaction'].lower() == 'true'
common_directory    = params['Common_Files_Directory']
# Add common utils directory to path
if common_directory not in sys.path:
    sys.path.append(common_directory)

from fileFinder import file_selector   # type: ignore
from debugOut import debug             # type: ignore
from get_bearer_token import getSignedJWT_fromfile, getBearerToken  #type: ignore

def get_bearer_token_global(fullCredsFile):
    # Generate BEARER_TOKEN only once and reuse it globally.
    global BEARER_TOKEN
    if BEARER_TOKEN is None:
        signedJWT, creds, creds["clientID"] = getSignedJWT_fromfile(fullCredsFile)
        BEARER_TOKEN = getBearerToken(signedJWT, creds)
    return BEARER_TOKEN

def getFileSetEnv(params):
    # Setup environment and get JWT.
    output_dir = "output"
    fileName, fullFile = file_selector('input ', current_directory)
    print(f"Input File: \n {fileName}   successfully uploaded...")
    credsFile, fullCredsFile = file_selector('Skyflow Credentials', current_directory)
    print(f"Credentials file: \n {credsFile} successfully uploaded...")
    #get type from the file name
    name, extension = os.path.splitext(fileName)
    file_type = extension[1:].lower()     #get file_type

    if file_type not in DATA_TYPES:       #not picking a supported data file
        print(f"ERROR: File type {file_type} not currently supported")
        print("Supported files are: mp3, wav, pdf, txt, csv, json, jpg, jpeg, tif, tiff, png, bmp")
        exit(1)
    else:                           #set payload api-options for file_type
        if file_type in params['ImageFiles']:
            payloadOption = params['image']
            group = 'image'
        elif file_type in params['AudioFiles']:
            payloadOption = params['audio']
            group = 'audio'
        elif file_type == 'pdf':
            payloadOption = params['pdf']
            group = 'pdf'
        else:
            payloadOption =  None      #detect api no options currently available
            group = None

    BEARER_TOKEN = get_bearer_token_global(fullCredsFile)
    return fullFile, current_directory, BEARER_TOKEN, payloadOption, file_type, name, group

def convert_file_to_base64(file_path):
    with open(file_path, "rb") as file:
        base64_encoded = base64.b64encode(file.read()).decode('utf-8')
    return base64_encoded

def save_base64_to_file(base64_string, output_path):
    with open(output_path, "wb") as file:
        file.write(base64.b64decode(base64_string))

def detect_file(file_path, payloadOption, file_type, fname, group):
    audioOut = None
    base64_file = convert_file_to_base64(file_path)
    
    url = VAULT_URL + '/' + 'v1/detect/file' 
    
    headers = {
        'Content-Type': 'application/json',
        'x-skyflow-account-id': SKYFLOW_ACCOUNT_ID,
        'Authorization': f'Bearer {BEARER_TOKEN}',
    }

    payload = {
        "file": base64_file,
        "data_format": file_type,
        "input_type": "BASE64",
        "vault_id": VAULT_ID,
        "restrict_entity_types": params['restrict_entity_types']
    }

    if payloadOption is not None:       #add api options to payload when available
        payload[group] = payloadOption
        if payloadOption.get("output_processed_audio") is not None:
            audioOut = payloadOption.get("output_processed_audio")

    response = requests.post(url, headers=headers, data=json.dumps(payload))

    if response.status_code == 200:
        return response.json(), audioOut
    else:
        print(f"Error {response.status_code}: {response.text}")  # Add more error context
        return {"error": response.status_code, "message": response.text}

def check_status(status_id):
    url = f'{VAULT_URL}/v1/detect/status/{status_id}?vault_id={VAULT_ID}'
    
    headers = {
        'Authorization': f'Bearer {BEARER_TOKEN}'
    }
    
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        return response.json()
    else:
        return {"error": response.status_code, "message": response.text}

def extract_base64_from_response(status_response):
    acceptable_types = ["redacted_transcription", "redacted_file", "redacted_image"]
    for output in status_response["output"]:
        if output.get("processedFileType") in acceptable_types:
            return output["processedFile"]
    return None

def getDateTime():
    now = datetime.now()
    return str(now.strftime("%m%d%Y%H%M%S"))

def replace_text_in_docx(doc, output_path):
    # Define a regex pattern to match text within square brackets
    pattern = re.compile(r'\[.*?\]')

    # Iterate through paragraphs and replace the text
    for para in doc.paragraphs:
        if pattern.search(para.text):
            replaced_text = pattern.sub(lambda match: '*' * len(match.group()), para.text)
            para.text = replaced_text

    # Iterate through tables (if the .docx file has tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if pattern.search(cell.text):
                    replaced_text = pattern.sub(lambda match: '*' * len(match.group()), cell.text)
                    cell.text = replaced_text

    # Save the modified document to the output path
    doc.save(output_path)
    print(f"\nRedacted DOCX file saved to: \n{output_path}")

def replace_text_in_txt(file_path, output_path):
    # Read the text from the file
    with open(file_path, 'r') as file:
        content = file.read()

    # Define a regex pattern to match text within square brackets
    pattern = re.compile(r'\[.*?\]')

    # Replace matched patterns with equivalent length '*' characters
    replaced_content = pattern.sub(lambda match: '*' * len(match.group()), content)

    # Write modified content back to new file
    with open(output_path, 'w') as file:
        file.write(replaced_content)
    print(f"\nSUCCESS: Redacted {file_type} saved to: \n{output_path}")

def redact_file(file_path, output_path, file_type):
    # Check it's a docx or txt file
    if file_type.lower() == 'docx':
        # Process the .docx file
        doc = docx.Document(file_path)
        replace_text_in_docx(doc, output_path)

    elif file_type.lower() == 'txt':
        # Process the .txt file
        replace_text_in_txt(file_path, output_path)

    else:
        print("Unsupported file format. Please provide a .docx or .txt file.")


def save_status_response(status_response, output_dir, file_type, fname):
    entities_processed_file = None
    for output in status_response.get('output', []):
        if output.get('processedFileType') == 'entities':
            entities_processed_file = output.get('processedFile')
            break

    if entities_processed_file:  # Decode base64-encoded content & save to file
        decoded_text = base64.b64decode(entities_processed_file).decode('utf-8')

        try:
            decoded_json = json.loads(decoded_text)
            text_file_path = os.path.join(output_dir, f"entities_{fname}-{file_type}_{getDateTime()}.json")
            with open(text_file_path, "w") as text_file:
                json.dump(decoded_json, text_file, indent=4)
            print(f"Decoded and formatted JSON text saved to: {text_file_path}")
        except json.JSONDecodeError:
            print("Error: Decoded text is not valid JSON.")
    else:
        print("Error: 'entities' processedFile not found in status response.")


def check_and_save_status(status_url_full, status_id, output_dir, file_type, fname, group, audioOut, max_attempts):
    # Check status of the file processing and save results.
    attempt = 0
    while attempt < max_attempts:
        status_response = check_status(status_id)

        outExt = None
        if audioOut is False and group == 'audio':
            outExt = 'txt'
        elif audioOut is True and group == 'audio':
            outExt = 'mp3'
        else:
            outExt = file_type

        if status_response.get("status") == "SUCCESS":
            processed_file_base64 = extract_base64_from_response(status_response)
            if processed_file_base64:
                timestamp = getDateTime()
                output_file_path = os.path.join(output_dir, f"tokenized_{fname}-{file_type}-{timestamp}.{outExt}")
                save_base64_to_file(processed_file_base64, output_file_path)
                print(f"\nSUCCESS: Processed {file_type} saved to: \n{output_file_path}")

                # Save the status response for debugging purposes
                save_status_response(status_response, output_dir, file_type, fname)

                # **Perform redaction if setting directive is set**
                # print(f"FULL_REDACTION: {FULL_REDACTION} (type: {type(FULL_REDACTION)})")

                if FULL_REDACTION == True:
                    redacted_output_path = os.path.join(output_dir, f"redacted_{fname}-{file_type}-{getDateTime()}.{file_type}")
                    redact_file(output_file_path, redacted_output_path, file_type)

            else:
                print("Processed file not found in the response.")
            break
        elif status_response.get("status") == "FAILED":
            print(f"Error in processing: {status_response.get('message')}")
            break
        else:
            attempt += 1
            print(f"Processing is still in progress (Attempt {attempt}/{max_attempts}), checking again in 2 seconds...")
            time.sleep(2)
    else:
        print("Maximum attempts reached. Process did not complete successfully.")

def run_files_detect(input_file_path, output_dir, payloadOption, file_type, fname, group, max_attempts=MAX_ATTEMPTS):
    os.makedirs(output_dir, exist_ok=True)

    # Detect file and get status URL
    detect_response, audioOut = detect_file(input_file_path, payloadOption, file_type, fname, group)

    # Check if response is a string or JSON (dict)
    if isinstance(detect_response, dict):
        if 'status_url' in detect_response:
            status_url_full = detect_response['status_url']
            status_id = status_url_full.split('/')[-1]  # Extract the unique identifier
            check_and_save_status(status_url_full, status_id, output_dir, file_type, fname, group, audioOut, max_attempts)
        else:
            print(f"Error in initial request: {detect_response.get('message')}")
    else:
        print(f"Error in initial request: {detect_response}")

if __name__ == "__main__":
    input_file_path, current_directory, BEARER_TOKEN, payloadOption, file_type, fname, group = getFileSetEnv(params)
    output_dir = current_directory + "/" + "output"
    run_files_detect(input_file_path, output_dir, payloadOption, file_type, fname, group, max_attempts=MAX_ATTEMPTS)